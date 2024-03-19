from docx import Document

def analyze_word_file(file_path):
    try:
        document = Document(file_path)

        if not document.paragraphs:
            return "Error: The document has no paragraphs."

        for paragraph in document.paragraphs:
            if not paragraph.text.strip():
                return "Error: Empty paragraph found in the document."

        return "No errors found."

    except Exception as e:
        return f"Error while analyzing the Word file: {e}"

if __name__ == "__main__":
    word_file_path = "anmol.docx"
    analysis_result = analyze_word_file(word_file_path)
    print(analysis_result)